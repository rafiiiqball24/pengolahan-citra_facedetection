from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "paper_deteksi_emosi_dua_kolom.docx"
FIGURE_INPUT = ROOT / "data" / "test.png"
RENDER_DIR = ROOT / "paper_rendered"
FIGURE_ANNOTATED = RENDER_DIR / "hasil_deteksi_test.png"

FONT = "Times New Roman"
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(80, 80, 80)
HEADER_FILL = "EDEDED"

DETECTIONS = [
    {"box": {"x": 1084, "y": 52, "width": 148, "height": 148}, "emotion": "sad", "score": 0.9447532892227173},
    {"box": {"x": 177, "y": 42, "width": 211, "height": 211}, "emotion": "happy", "score": 0.9989953637123108},
    {"box": {"x": 633, "y": 88, "width": 171, "height": 171}, "emotion": "happy", "score": 0.843101441860199},
]


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=8.3, fill=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_two_columns(section):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def add_paragraph(doc, text="", size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title.upper())
    set_run_font(run, size=10, bold=True)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=8, italic=True, color=MUTED)
    return paragraph


def add_reference(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=8)


def build_annotated_figure():
    if not FIGURE_INPUT.exists():
        return None

    RENDER_DIR.mkdir(exist_ok=True)
    image = Image.open(FIGURE_INPUT).convert("RGB")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 28)
    except OSError:
        font = ImageFont.load_default()

    colors = {
        "happy": (234, 179, 8),
        "sad": (59, 130, 246),
        "angry": (239, 68, 68),
        "surprise": (168, 85, 247),
        "neutral": (34, 197, 94),
        "fear": (14, 165, 233),
        "disgust": (16, 185, 129),
    }
    for detection in DETECTIONS:
        box = detection["box"]
        x, y, width, height = box["x"], box["y"], box["width"], box["height"]
        emotion = detection["emotion"]
        score = detection["score"]
        color = colors.get(emotion, (34, 197, 94))
        draw.rectangle((x, y, x + width, y + height), outline=color, width=5)
        label = f"{emotion} {score:.2f}"
        bbox = draw.textbbox((x, y), label, font=font)
        label_width = bbox[2] - bbox[0]
        label_height = bbox[3] - bbox[1]
        label_top = max(0, y - label_height - 10)
        draw.rectangle((x, label_top, x + label_width + 12, label_top + label_height + 8), fill=color)
        draw.text((x + 6, label_top + 3), label, fill=(255, 255, 255), font=font)

    image.save(FIGURE_ANNOTATED)
    return FIGURE_ANNOTATED


def configure_document(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(4)


def add_front_matter(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(
        "Deteksi Emosi Wajah Menggunakan Haar Cascade dan CNN Mini-Xception "
        "Berbasis FER2013"
    )
    set_run_font(run, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    run = meta.add_run("Nama Penulis: ____________________    Program Studi: ____________________")
    set_run_font(run, size=10)

    abstract = doc.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_after = Pt(4)
    run = abstract.add_run("Abstrak - ")
    set_run_font(run, size=9, bold=True)
    run = abstract.add_run(
        "Paper ini membahas implementasi sistem deteksi emosi wajah berdasarkan "
        "project Python yang memadukan OpenCV Haar Cascade untuk deteksi wajah dan "
        "model CNN Mini-Xception terlatih pada FER2013 untuk klasifikasi emosi. "
        "Sistem menerima input berupa gambar, video, atau webcam; mengekstraksi wajah; "
        "mengubah region wajah ke grayscale; melakukan resize sesuai ukuran input "
        "model; menormalisasi piksel; lalu menghasilkan label emosi dan skor prediksi. "
        "Pengujian lokal pada file data/test.png mendeteksi tiga wajah dengan emosi "
        "dominan happy. Hasil tersebut menunjukkan bahwa pipeline ringan berbasis "
        "Haar Cascade dan CNN dapat digunakan sebagai prototipe real-time, meskipun "
        "kinerja sistem masih dipengaruhi pencahayaan, pose wajah, kualitas citra, "
        "dan keterbatasan distribusi data FER2013."
    )
    set_run_font(run, size=9)

    keywords = doc.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.space_after = Pt(6)
    run = keywords.add_run("Kata kunci - ")
    set_run_font(run, size=9, bold=True)
    run = keywords.add_run("deteksi emosi, ekspresi wajah, CNN, Mini-Xception, FER2013, OpenCV")
    set_run_font(run, size=9)


def add_component_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Komponen", bold=True, fill=HEADER_FILL)
    set_cell_text(table.rows[0].cells[1], "Implementasi pada Project", bold=True, fill=HEADER_FILL)

    rows = [
        ("Input", "Gambar, video, atau webcam melalui OpenCV dan GUI CustomTkinter."),
        ("Deteksi wajah", "OpenCV CascadeClassifier dengan Haar Cascade frontal face."),
        ("Model emosi", "fer2013_mini_XCEPTION.102-0.66.hdf5, tujuh kelas FER2013."),
        ("Preprocessing", "Grayscale, resize, normalisasi [0,1], lalu pemetaan ke [-1,1]."),
        ("Output", "Bounding box, label emosi, skor probabilitas, dan emosi dominan."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)
    add_caption(doc, "Tabel 1. Komponen utama sistem deteksi emosi wajah.")


def main():
    annotated = build_annotated_figure()
    doc = Document()
    configure_document(doc)
    add_front_matter(doc)

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_two_columns(section)

    add_heading(doc, "I. Pendahuluan")
    add_paragraph(
        doc,
        "Ekspresi wajah merupakan bentuk komunikasi nonverbal yang dapat "
        "menggambarkan keadaan emosional manusia. Pada aplikasi interaksi "
        "manusia-komputer, pembelajaran daring, robotika sosial, dan sistem "
        "monitoring, kemampuan membaca ekspresi wajah dapat menjadi komponen "
        "pendukung pengambilan keputusan. Project emotion-detection pada repository "
        "ini mengimplementasikan pendekatan praktis untuk mengenali ekspresi wajah "
        "dengan memanfaatkan model yang sudah dilatih, sehingga fokus proyek berada "
        "pada integrasi pipeline inferensi dan antarmuka penggunaan."
    )
    add_paragraph(
        doc,
        "Secara umum, sistem memiliki dua tahap utama: deteksi lokasi wajah dan "
        "klasifikasi emosi pada area wajah yang terdeteksi. Tahap pertama memakai "
        "Haar Cascade dari OpenCV, sedangkan tahap kedua memakai model CNN "
        "Mini-Xception berbasis dataset FER2013. Pendekatan ini dipilih karena "
        "relatif ringan, mudah dijalankan di komputer lokal, dan cukup sesuai untuk "
        "prototipe berbasis citra, video, maupun webcam."
    )

    add_heading(doc, "II. Tinjauan Pustaka")
    add_paragraph(
        doc,
        "FER2013 dikenal sebagai benchmark klasifikasi ekspresi wajah yang berasal "
        "dari Challenge in Representation Learning. Dataset ini umum digunakan untuk "
        "tujuh kelas emosi, yaitu angry, disgust, fear, happy, sad, surprise, dan "
        "neutral. Pada project ini, label kelas diambil dari fungsi get_labels pada "
        "utils/datasets.py, sedangkan bobot model emosi diambil dari file HDF5 yang "
        "tersedia pada folder trained_models/emotion_models."
    )
    add_paragraph(
        doc,
        "Untuk deteksi wajah, OpenCV menyediakan CascadeClassifier yang bekerja "
        "dengan cascade berbasis fitur Haar. Metode ini bukan pendekatan paling baru, "
        "tetapi masih sering dipakai untuk prototipe karena sederhana dan cepat. "
        "Untuk klasifikasi, arsitektur Xception memperkenalkan depthwise separable "
        "convolution yang lebih hemat parameter dibanding konvolusi standar. Model "
        "Mini-Xception pada project ini mengadopsi gagasan tersebut melalui blok "
        "separable convolution, batch normalization, residual connection, global "
        "average pooling, dan softmax."
    )

    add_heading(doc, "III. Metodologi")
    add_paragraph(
        doc,
        "Analisis metodologi dilakukan dengan membaca main_emotion_classifier.py, "
        "video_emotion_color_demo.py, gui_emotion_app.py, utils/inference.py, "
        "utils/preprocessor.py, utils/datasets.py, dan models/cnn.py. Pada mode file, "
        "program menerima path input. Jika input berupa gambar, program langsung "
        "memproses satu frame. Jika input berupa video, program mengekstraksi frame "
        "ke direktori sementara, memproses setiap frame, lalu menghapus direktori "
        "sementara setelah inferensi selesai."
    )
    add_paragraph(
        doc,
        "Setiap frame diproses dengan langkah berikut: citra dikonversi ke grayscale, "
        "wajah dideteksi menggunakan CascadeClassifier, koordinat wajah diberi offset, "
        "area wajah dipotong dan di-resize sesuai input model, nilai piksel "
        "dinormalisasi, dimensi batch dan channel ditambahkan, lalu classifier "
        "menghasilkan probabilitas untuk tujuh kelas emosi. Label akhir dipilih "
        "berdasarkan nilai probabilitas tertinggi. Untuk video dan webcam, sistem "
        "juga menghitung emosi dominan berdasarkan kemunculan label pada frame."
    )
    add_component_table(doc)
    add_paragraph(
        doc,
        "Arsitektur Mini-Xception pada models/cnn.py tersusun dari blok awal "
        "Conv2D dan beberapa modul separable convolution. Setiap modul menggunakan "
        "jalur residual 1x1 convolution, dua lapis SeparableConv2D, BatchNormalization, "
        "aktivasi ReLU, dan MaxPooling2D. Bagian akhir menggunakan Conv2D sejumlah "
        "kelas, GlobalAveragePooling2D, dan aktivasi softmax bernama predictions."
    )

    add_heading(doc, "IV. Hasil dan Pembahasan")
    add_paragraph(
        doc,
        "Pengujian dilakukan pada file data/test.png menggunakan script "
        "main_emotion_classifier.py. Program mendeteksi tiga wajah. Wajah pertama "
        "terdeteksi pada koordinat x=1084, y=52, w=148, h=148 dan diprediksi sad "
        "dengan skor 0,9448. Wajah kedua terdeteksi pada x=177, y=42, w=211, h=211 "
        "dan diprediksi happy dengan skor 0,9990. Wajah ketiga terdeteksi pada "
        "x=633, y=88, w=171, h=171 dan diprediksi happy dengan skor 0,8431. Fungsi "
        "agregasi most_frequent menghasilkan emosi dominan happy karena label happy "
        "muncul paling banyak."
    )
    if annotated and annotated.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run().add_picture(str(annotated), width=Cm(7.4))
        add_caption(doc, "Gambar 1. Hasil anotasi deteksi wajah dan klasifikasi emosi pada data/test.png.")
    elif FIGURE_INPUT.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run().add_picture(str(FIGURE_INPUT), width=Cm(7.4))
        add_caption(doc, "Gambar 1. Citra uji lokal yang digunakan pada evaluasi.")
    add_paragraph(
        doc,
        "Hasil tersebut memperlihatkan bahwa pipeline end-to-end sudah berjalan: "
        "detector menghasilkan koordinat wajah, classifier menghasilkan label dan "
        "skor prediksi, serta fungsi agregasi mengembalikan emosi dominan. Namun, "
        "pengujian satu citra belum cukup untuk menyimpulkan akurasi sistem secara "
        "menyeluruh. Evaluasi yang lebih kuat perlu memakai data berlabel dalam "
        "jumlah lebih besar, confusion matrix, precision, recall, F1-score, serta "
        "pengujian pada kondisi pencahayaan dan pose yang beragam."
    )
    add_paragraph(
        doc,
        "Pada implementasi real-time, video_emotion_color_demo.py menampilkan "
        "bounding box, label emosi, probabilitas, dan daftar skor tiap kelas pada "
        "frame webcam. GUI pada gui_emotion_app.py memperluas fungsi tersebut dengan "
        "pilihan open image, open video, dan start webcam. GUI juga menampilkan "
        "ringkasan source, dominant emotion, jumlah wajah, frame, dan probabilitas "
        "kelas sehingga sistem lebih mudah digunakan oleh pengguna non-teknis."
    )

    add_heading(doc, "V. Kesimpulan")
    add_paragraph(
        doc,
        "Project emotion-detection berhasil menggabungkan deteksi wajah berbasis Haar "
        "Cascade dan klasifikasi emosi berbasis CNN Mini-Xception. Sistem dapat "
        "memproses gambar, video, dan webcam, lalu memberikan bounding box, label "
        "emosi, skor probabilitas, dan emosi dominan. Keunggulan utama pendekatan ini "
        "adalah sederhana, ringan, dan mudah dijalankan sebagai prototipe. "
        "Keterbatasannya adalah sensitivitas terhadap kualitas deteksi wajah, "
        "pencahayaan, pose, occlusion, serta bias dataset FER2013. Pengembangan "
        "berikutnya dapat diarahkan pada evaluasi kuantitatif yang lebih lengkap, "
        "augmentasi data lokal, penggunaan face detector modern, serta deployment "
        "aplikasi yang lebih stabil."
    )

    add_heading(doc, "Daftar Pustaka")
    references = [
        "[1] I. J. Goodfellow et al., \"Challenges in Representation Learning: A report on three machine learning contests,\" arXiv:1307.0414, 2013. URL: https://arxiv.org/abs/1307.0414",
        "[2] OpenCV, \"Cascade Classifier,\" OpenCV Documentation. URL: https://docs.opencv.org/4.x/db/d28/tutorial_cascade_classifier.html",
        "[3] F. Chollet, \"Xception: Deep Learning with Depthwise Separable Convolutions,\" arXiv:1610.02357, 2016. URL: https://arxiv.org/abs/1610.02357",
        "[4] O. Arriaga, \"face_classification: Real-time face detection and emotion/gender classification using FER2013/IMDB datasets with a Keras CNN model and OpenCV,\" GitHub repository. URL: https://github.com/oarriaga/face_classification",
        "[5] O. Arriaga, P. G. Ploger, and M. Valdenegro-Toro, \"Real-time Convolutional Neural Networks for Emotion and Gender Classification,\" 2017. URL: https://github.com/oarriaga/face_classification/blob/master/report.pdf",
    ]
    for reference in references:
        add_reference(doc, reference)

    doc.core_properties.title = "Deteksi Emosi Wajah Menggunakan Haar Cascade dan CNN Mini-Xception"
    doc.core_properties.subject = "Paper dua kolom berbasis project emotion-detection"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "emotion detection, FER2013, OpenCV, CNN, Mini-Xception"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
